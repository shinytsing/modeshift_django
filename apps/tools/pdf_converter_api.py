#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转换引擎API
支持PDF与Word、图片等格式的相互转换
"""

import json
import logging
import os

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

try:
    import fitz  # PyMuPDF

    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    print("警告: PyMuPDF (fitz) 未安装，PDF转换功能将受限")

try:
    import pdfplumber
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    print("警告: pypdf/pdfplumber 未安装，PDF转换功能将受限")
import base64
import io
import uuid

from PIL import Image

# 配置日志
logger = logging.getLogger(__name__)


class PDFConverter:
    """PDF转换引擎核心类"""

    def __init__(self):
        self.supported_formats = {
            "pdf": [".pdf"],
            "word": [".doc", ".docx"],
            "image": [".jpg", ".jpeg", ".png", ".bmp", ".tiff"],
            "text": [".txt"],
        }

    def validate_file(self, file, expected_type):
        """验证文件格式"""
        if not file:
            return False, "文件不能为空"

        # 检查file对象是否有name属性
        if not hasattr(file, "name") or not file.name:
            return False, "无效的文件对象"

        file_ext = os.path.splitext(file.name)[1].lower()

        # 检查文件大小 (限制为50MB)
        if hasattr(file, "size") and file.size > 50 * 1024 * 1024:
            return False, "文件大小不能超过50MB"

        # 检查文件格式兼容性
        if expected_type in self.supported_formats:
            if file_ext not in self.supported_formats[expected_type]:
                # 提供智能提示和自动切换建议
                suggestion = self._get_conversion_suggestion(file_ext, expected_type)
                return False, suggestion

        return True, "文件验证通过"

    def _get_conversion_suggestion(self, file_ext, current_type):
        """获取转换类型建议"""
        # 检查file_ext是否为字符串
        if not isinstance(file_ext, str):
            return f"无效的文件扩展名: {file_ext}"

        # 定义文件类型到转换类型的映射
        file_type_mapping = {
            ".pdf": "pdf",
            ".doc": "word",
            ".docx": "word",
            ".jpg": "image",
            ".jpeg": "image",
            ".png": "image",
            ".bmp": "image",
            ".tiff": "image",
            ".gif": "image",
            ".txt": "text",
        }

        # 定义转换类型到操作类型的映射
        conversion_mapping = {
            "pdf": {"pdf-to-word": "PDF转Word", "pdf-to-image": "PDF转图片"},
            "word": {"word-to-pdf": "Word转PDF"},
            "image": {"image-to-pdf": "图片转PDF"},
            "text": {"text-to-pdf": "文本转PDF"},
        }

        # 获取文件的实际类型
        actual_type = file_type_mapping.get(file_ext, "unknown")

        if actual_type == "unknown":
            return f"不支持的文件格式: {file_ext}。请使用支持的文件格式。"

        # 获取当前转换类型的显示名称
        current_display = conversion_mapping.get(actual_type, {}).get(current_type, current_type)

        # 获取建议的转换类型
        suggested_conversions = conversion_mapping.get(actual_type, {})

        if not suggested_conversions:
            return f"文件格式 {file_ext} 不支持任何转换操作。"

        # 构建建议信息
        suggestion = f"文件格式 {file_ext} 与当前转换类型 '{current_display}' 不兼容。\n\n"
        suggestion += "建议的转换类型：\n"

        for conv_type, display_name in suggested_conversions.items():
            suggestion += f"• {display_name} ({conv_type})\n"

        suggestion += f"\n请切换到适合的转换类型，或上传 {actual_type} 格式的文件。"

        return suggestion

    def pdf_to_word(self, pdf_file):
        """PDF转Word - 真实实现"""
        try:
            # 检查pdf2docx库是否可用
            try:
                from pdf2docx import Converter
            except ImportError:
                return False, "pdf2docx库未安装，无法进行PDF转Word转换", None

            # 重置文件指针
            pdf_file.seek(0)

            # 使用pdf2docx进行真实转换
            import os
            import tempfile

            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
                temp_pdf.write(pdf_file.read())
                temp_pdf_path = temp_pdf.name

            # 创建临时输出文件路径
            temp_docx_path = temp_pdf_path.replace(".pdf", ".docx")

            try:
                # 检查输入PDF文件是否存在
                if not os.path.exists(temp_pdf_path):
                    return False, "临时PDF文件创建失败", None

                # 使用pdf2docx进行转换，改进页面布局处理
                cv = Converter(temp_pdf_path)

                try:
                    # 改进转换参数，优化页面布局，并添加错误处理
                    import logging

                    # 临时抑制pdf2docx的WARNING和ERROR日志，只保留CRITICAL
                    pdf2docx_logger = logging.getLogger("pdf2docx")
                    original_level = pdf2docx_logger.getEffectiveLevel()
                    pdf2docx_logger.setLevel(logging.CRITICAL)

                    try:
                        cv.convert(temp_docx_path, start=0, end=None, pages=None, zoom_x=1.0, zoom_y=1.0, crop=(0, 0, 0, 0))
                    finally:
                        # 恢复原始日志级别
                        pdf2docx_logger.setLevel(original_level)

                except Exception as conv_error:
                    cv.close()
                    logger.warning(f"pdf2docx转换时出现警告（但可能仍然成功）: {conv_error}")
                    # 继续检查是否产生了输出文件

                cv.close()

                # 检查输出文件是否存在
                if not os.path.exists(temp_docx_path):
                    return False, "转换失败：输出Word文件未生成", None

                # 读取转换后的文件
                with open(temp_docx_path, "rb") as docx_file:
                    docx_content = docx_file.read()

                if len(docx_content) == 0:
                    return False, "转换后的文件为空，可能是扫描版PDF或内容无法识别", None

                # 检查转换结果是否包含实际内容，并设置中文字体
                try:
                    from docx import Document

                    doc = Document(io.BytesIO(docx_content))

                    # 设置文档的中文字体
                    from docx.oxml.shared import qn

                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "宋体"
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

                    text_content = ""
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"

                    logger.info(f"pdf2docx提取的文本长度: {len(text_content.strip())}")

                    # 检查内容质量 - 如果内容太少、包含乱码或不合理，切换到OCR
                    content_length = len(text_content.strip())

                    # 检查是否包含乱码或不合理内容
                    suspicious_patterns = ["BERBER", "x1", "23450m", "ERROR"]
                    contains_suspicious = any(pattern in text_content for pattern in suspicious_patterns)

                    # 检查是否主要是数字或特殊字符（可能是提取错误）
                    non_space_chars = text_content.replace(" ", "").replace("\n", "")
                    if non_space_chars:
                        alpha_ratio = sum(c.isalpha() for c in non_space_chars) / len(non_space_chars)
                    else:
                        alpha_ratio = 0

                    logger.info(
                        f"内容质量检查: 长度={content_length}, 包含可疑内容={contains_suspicious}, 字母比例={alpha_ratio:.2f}"
                    )

                    # 如果内容太少、包含可疑内容或字母比例太低，切换到OCR
                    if content_length < 50 or contains_suspicious or alpha_ratio < 0.3:
                        logger.info("pdf2docx提取内容不足，切换到OCR模式")
                        # 自动启用 OCR 降级为扫描件识别
                        ocr_success, ocr_result, ocr_type = self._ocr_pdf_to_word_from_path(temp_pdf_path)
                        if ocr_success:
                            logger.info("OCR转换成功，使用OCR结果")
                            # 清理临时文件
                            try:
                                os.unlink(temp_pdf_path)
                                if os.path.exists(temp_docx_path):
                                    os.unlink(temp_docx_path)
                            except Exception:
                                pass
                            return True, ocr_result, "pdf_to_word"
                        else:
                            logger.warning(f"OCR转换也失败: {ocr_result}")
                            # 如果OCR也失败，但pdf2docx有输出，继续使用pdf2docx结果
                            if len(docx_content) > 1000:  # 文件不为空
                                logger.info("使用pdf2docx结果，尽管内容较少")
                            else:
                                # 清理临时文件
                                try:
                                    os.unlink(temp_pdf_path)
                                    if os.path.exists(temp_docx_path):
                                        os.unlink(temp_docx_path)
                                except Exception:
                                    pass
                                return False, f"转换失败: pdf2docx提取内容不足，OCR也失败: {ocr_result}", None
                    else:
                        logger.info("pdf2docx提取内容充足，使用pdf2docx结果")

                    # 保存修改后的文档（包含中文字体设置）
                    import io as _bio

                    doc_buffer = _bio.BytesIO()
                    doc.save(doc_buffer)
                    docx_content = doc_buffer.getvalue()

                except Exception as check_error:
                    logger.warning(f"转换结果检查失败: {check_error}")
                    # 如果检查失败但有输出文件，尝试使用OCR
                    if len(docx_content) < 1000:
                        logger.info("由于检查失败且文件较小，尝试OCR转换")
                        try:
                            ocr_success, ocr_result, ocr_type = self._ocr_pdf_to_word_from_path(temp_pdf_path)
                            if ocr_success:
                                # 清理临时文件
                                try:
                                    os.unlink(temp_pdf_path)
                                    if os.path.exists(temp_docx_path):
                                        os.unlink(temp_docx_path)
                                except Exception:
                                    pass
                                return True, ocr_result, "pdf_to_word"
                        except Exception as ocr_error:
                            logger.warning(f"OCR备用转换失败: {ocr_error}")
                    # 继续处理，不因为检查失败而中断

                # 清理临时文件
                try:
                    os.unlink(temp_pdf_path)
                    os.unlink(temp_docx_path)
                except Exception:
                    pass

                return True, docx_content, "pdf_to_word"

            except Exception as conversion_error:
                # 清理临时文件
                try:
                    if os.path.exists(temp_pdf_path):
                        os.unlink(temp_pdf_path)
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
                except Exception:
                    pass

                # 提供更详细的错误信息
                error_msg = str(conversion_error)
                if "Package not found" in error_msg:
                    return False, "PDF文件损坏或格式不支持，请检查文件完整性", None
                elif "Permission denied" in error_msg:
                    return False, "文件访问权限不足，请检查文件权限", None
                else:
                    return False, f"PDF转Word转换失败: {error_msg}", None

        except Exception as e:
            logger.error(f"PDF转Word失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def _ocr_pdf_to_word_from_path(self, pdf_path: str):
        """对扫描版PDF执行OCR识别并输出Word（docx字节）
        依赖: pytesseract + 系统 tesseract 二进制
        """
        try:
            try:
                import pytesseract
                from docx import Document
                from docx.shared import Pt
            except ImportError:
                return False, "OCR依赖未安装：需要 pytesseract（以及系统 tesseract）。请安装后重试。", None

            # 检查系统 tesseract 是否可用
            try:
                _ = pytesseract.get_tesseract_version()
            except Exception:
                return False, "系统未安装 tesseract 或未在PATH中，无法进行OCR。请安装 tesseract 后重试。", None

            # 打开PDF并逐页渲染为图片
            if not FITZ_AVAILABLE:
                return False, "PyMuPDF未安装，无法进行OCR渲染", None

            doc = fitz.open(pdf_path)
            ocr_texts = []
            # 使用较高的DPI 提升识别率
            zoom = 300 / 72.0
            mat = fitz.Matrix(zoom, zoom)

            import io as _io

            from PIL import Image as PILImage  # 避免命名冲突

            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                img = PILImage.open(_io.BytesIO(img_bytes))

                # 语言优先中文+英文（若未安装对应语言，pytesseract会回退或报错）
                try:
                    # 使用中文+英文OCR，移除字符白名单限制以支持更多中文字符
                    config = "--oem 3 --psm 6"
                    text = pytesseract.image_to_string(img, lang="chi_sim+eng", config=config)
                except Exception as ocr_error:
                    logger.warning(f"中文OCR失败，回退到英文: {ocr_error}")
                    try:
                        # 回退英文
                        text = pytesseract.image_to_string(img, lang="eng")
                    except Exception as eng_error:
                        logger.warning(f"英文OCR也失败: {eng_error}")
                        text = ""

                text = text.strip()
                if text:
                    ocr_texts.append(text)
                else:
                    ocr_texts.append("")

            doc.close()

            # 生成docx，改进的页面结构处理 - 优化版本
            document = Document()

            # 设置文档的默认字体和编码
            style = document.styles["Normal"]
            font = style.font
            font.name = "宋体"  # 中文字体
            font.size = Pt(12)

            # 设置中文字体备选方案
            from docx.oxml.shared import qn

            try:
                font.element.set(qn("w:eastAsia"), "宋体")
                font.element.set(qn("w:ascii"), "Times New Roman")
                font.element.set(qn("w:hAnsi"), "Times New Roman")
                font.element.set(qn("w:cs"), "宋体")
            except Exception as font_error:
                logger.warning(f"设置字体时出现警告: {font_error}")
                # 继续处理，不中断转换

                # 优化的文本处理：保持原始页面结构，减少不必要的分页
            all_text = ""
            for page_index, page_text in enumerate(ocr_texts):
                if page_text.strip():
                    if page_index > 0 and all_text:
                        # 检查页面连接性，如果内容较短或看起来连续，直接连接
                        last_char = all_text.strip()[-1] if all_text.strip() else ""
                        first_char = page_text.strip()[0] if page_text.strip() else ""

                        # 智能连接：如果是连续内容或短页面，使用空格连接
                        # 改进中文文本连接逻辑
                        if (last_char.isalnum() and first_char.islower()) or len(page_text.strip()) < 200:
                            # 中文之间不需要空格，英文之间需要空格
                            # 修复字符编码检测逻辑
                            if last_char.isalpha() and first_char.isalpha():
                                # 如果都是中文字符，直接连接
                                if ord(last_char) > 127 and ord(first_char) > 127:
                                    all_text += page_text.strip()
                                else:
                                    all_text += " " + page_text.strip()
                            else:
                                all_text += " " + page_text.strip()
                        else:
                            # 否则使用简单的段落分隔
                            all_text += "\n\n" + page_text.strip()
                    else:
                        all_text += page_text.strip()

            # 简化的分段处理：保持原始结构
            if all_text:
                # 按自然段落分割
                paragraphs = []
                current_paragraph = ""

                # 按行处理，但避免过度分割
                lines = all_text.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        # 空行：结束当前段落
                        if current_paragraph:
                            paragraphs.append(current_paragraph)
                            current_paragraph = ""
                    else:
                        # 累积段落内容
                        if current_paragraph:
                            current_paragraph += " " + line
                        else:
                            current_paragraph = line

                # 添加最后一个段落
                if current_paragraph:
                    paragraphs.append(current_paragraph)

                # 将段落添加到文档中
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        # 直接添加段落，不再进行复杂的分割
                        p = document.add_paragraph()
                        p.add_run(paragraph_text.strip())

            import io as _bio

            buffer = _bio.BytesIO()
            document.save(buffer)
            return True, buffer.getvalue(), "pdf_to_word"

        except Exception as e:
            logger.error(f"OCR转换失败: {str(e)}")
            return False, f"OCR转换失败: {str(e)}", None

    def word_to_pdf(self, word_file):
        """Word转PDF - 真实实现"""
        try:
            # 重置文件指针
            word_file.seek(0)

            # 检查必要的库
            try:
                from docx import Document
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            except ImportError as e:
                error_msg = f"缺少必要的库: {str(e)}\\n"
                error_msg += "请安装以下依赖:\\n"
                error_msg += "1. python-docx: pip install python-docx\\n"
                error_msg += "2. reportlab: pip install reportlab\\n"
                error_msg += "3. 如果已安装，请重启服务器\\n"
                error_msg += "4. 检查Python环境是否正确"
                return False, error_msg, None

            # 创建临时文件
            import os
            import tempfile

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                temp_docx.write(word_file.read())
                temp_docx_path = temp_docx.name

            # 创建临时输出文件路径
            temp_docx_path.replace(".docx", ".pdf")

            try:
                # 检查临时文件是否存在
                if not os.path.exists(temp_docx_path):
                    return False, "临时Word文件创建失败", None

                # 读取Word文档
                doc = Document(temp_docx_path)

                # 创建PDF缓冲区
                pdf_buffer = io.BytesIO()

                # 创建PDF文档
                pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
                story = []

                # 获取样式
                styles = getSampleStyleSheet()
                normal_style = styles["Normal"]

                # 设置中文字体支持
                try:
                    # 使用reportlab内置的中文字体
                    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                    normal_style.fontName = "STSong-Light"
                    normal_style.fontSize = 12
                    normal_style.leading = 14
                    normal_style.alignment = 0  # 左对齐
                except Exception:
                    try:
                        # 尝试使用系统中文字体
                        import platform

                        if platform.system() == "Darwin":  # macOS
                            try:
                                pdfmetrics.registerFont(TTFont("PingFang", "/System/Library/Fonts/PingFang.ttc"))
                                normal_style.fontName = "PingFang"
                            except Exception:
                                pdfmetrics.registerFont(TTFont("HiraginoSans", "/System/Library/Fonts/STHeiti Light.ttc"))
                                normal_style.fontName = "HiraginoSans"
                        elif platform.system() == "Windows":
                            pdfmetrics.registerFont(TTFont("SimSun", "C:/Windows/Fonts/simsun.ttc"))
                            normal_style.fontName = "SimSun"
                        else:  # Linux
                            pdfmetrics.registerFont(TTFont("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
                            normal_style.fontName = "DejaVuSans"

                        normal_style.fontSize = 12
                        normal_style.leading = 14
                        normal_style.alignment = 0
                    except Exception as e2:
                        # 如果都不可用，使用默认字体
                        normal_style.fontName = "Helvetica"
                        normal_style.fontSize = 12
                        normal_style.leading = 14
                        normal_style.alignment = 0
                        logger.warning(f"无法加载中文字体，将使用默认字体: {str(e2)}")

                # 简化的图片处理 - 直接从文档关系中提取所有图片
                total_images_added = 0
                logger.info("开始处理Word文档中的图片...")

                # 直接从文档的关系部分提取所有图片
                if hasattr(doc.part, "related_parts"):
                    for rel_id, rel_part in doc.part.related_parts.items():
                        if hasattr(rel_part, "blob") and rel_part.blob:
                            try:
                                content_type = getattr(rel_part, "content_type", "")
                                # 检查是否是图片类型
                                if any(
                                    img_type in content_type.lower()
                                    for img_type in ["image", "jpeg", "png", "gif", "bmp", "tiff"]
                                ):
                                    img_buffer = io.BytesIO(rel_part.blob)
                                    try:
                                        from PIL import Image as PILImage
                                        from reportlab.platypus import Image as RLImage

                                        # 验证并打开图片
                                        pil_img = PILImage.open(img_buffer)
                                        img_width, img_height = pil_img.size

                                        # 计算合适的PDF尺寸
                                        max_width = 450
                                        max_height = 350
                                        ratio = min(max_width / img_width, max_height / img_height, 1.0)  # 不放大
                                        pdf_width = img_width * ratio
                                        pdf_height = img_height * ratio

                                        # 添加图片到PDF
                                        img_buffer.seek(0)
                                        img = RLImage(img_buffer, width=pdf_width, height=pdf_height)
                                        story.append(img)
                                        story.append(Spacer(1, 12))
                                        total_images_added += 1
                                        logger.info(
                                            f"成功添加图片 {total_images_added}: {rel_id}, 尺寸: {pdf_width:.1f}x{pdf_height:.1f}"
                                        )

                                    except Exception as img_err:
                                        logger.warning(f"处理图片 {rel_id} 失败: {img_err}")

                            except Exception as rel_err:
                                logger.debug(f"处理关系 {rel_id} 失败: {rel_err}")

                logger.info(f"图片处理完成，共添加 {total_images_added} 个图片")

                # 检查文档是否有内容
                if not doc.paragraphs and not doc.inline_shapes:
                    paragraph = Paragraph("空文档", normal_style)
                    story.append(paragraph)
                else:
                    # 处理段落和图片
                    for element in doc.element.body:
                        if element.tag.endswith("p"):  # 段落
                            paragraph = doc.paragraphs[
                                len([e for e in doc.element.body[: doc.element.body.index(element)] if e.tag.endswith("p")])
                            ]
                            if paragraph.text.strip():
                                # 创建段落
                                para = Paragraph(paragraph.text, normal_style)
                                story.append(para)
                                story.append(Spacer(1, 6))  # 添加间距
                            else:
                                # 空行
                                story.append(Spacer(1, 12))
                        elif element.tag.endswith("drawing"):  # 图片
                            try:
                                # 改进的图片提取逻辑
                                from reportlab.platypus import Image as RLImage

                                # 查找所有可能的图片元素 - 改进的检测算法
                                image_elements = []

                                # 方法1: 查找pic:pic元素（带命名空间）
                                pic_elements = element.findall(
                                    ".//pic:pic",
                                    namespaces={"pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"},
                                )
                                if pic_elements:
                                    image_elements.extend(pic_elements)

                                # 方法2: 查找无命名空间的pic元素
                                if not image_elements:
                                    pic_elements = element.findall(".//pic:pic")
                                    if pic_elements:
                                        image_elements.extend(pic_elements)

                                # 方法3: 查找所有可能的图片引用
                                if not image_elements:
                                    # 查找所有包含图片引用的元素
                                    for img_elem in element.iter():
                                        if "embed" in img_elem.attrib or "link" in img_elem.attrib:
                                            image_elements.append(img_elem)

                                # 方法4: 查找blip元素（带命名空间）
                                if not image_elements:
                                    blip_elements = element.findall(
                                        ".//a:blip", namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                                    )
                                    if blip_elements:
                                        image_elements.extend(blip_elements)

                                # 方法5: 查找无命名空间的blip元素
                                if not image_elements:
                                    blip_elements = element.findall(".//blip")
                                    if blip_elements:
                                        image_elements.extend(blip_elements)

                                # 方法6: 查找所有可能的图片关系引用
                                if not image_elements:
                                    for img_elem in element.iter():
                                        # 检查所有可能的图片引用属性
                                        for attr_name, attr_value in img_elem.attrib.items():
                                            if (
                                                "embed" in attr_name
                                                or "link" in attr_name
                                                or "r:embed" in attr_name
                                                or "r:link" in attr_name
                                            ):
                                                image_elements.append(img_elem)
                                                break

                                # 方法7: 直接从文档的关系中查找图片
                                if not image_elements and hasattr(doc.part, "related_parts"):
                                    # 如果没找到图片元素，但文档有关系部分，尝试直接处理
                                    for rel_id, rel_part in doc.part.related_parts.items():
                                        if hasattr(rel_part, "content_type") and "image" in rel_part.content_type:
                                            # 创建一个虚拟的图片元素
                                            from lxml.etree import Element

                                            virtual_elem = Element("virtual_image")
                                            virtual_elem.set("embed", rel_id)
                                            image_elements.append(virtual_elem)

                                # 处理找到的图片元素
                                for shape in image_elements:
                                    # 改进的图片引用获取算法
                                    rId = None

                                    # 尝试多种方式获取图片引用
                                    if "embed" in shape.attrib:
                                        rId = shape.get("embed")
                                    elif "link" in shape.attrib:
                                        rId = shape.get("link")
                                    else:
                                        # 查找子元素中的引用
                                        blip = shape.find(
                                            ".//a:blip",
                                            namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                                        )
                                        if not blip:
                                            blip = shape.find(".//a:blip")

                                        if blip is not None:
                                            rId = blip.get(
                                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                            )
                                            if not rId:
                                                rId = blip.get("embed")
                                            if not rId:
                                                rId = blip.get("link")

                                    # 如果还是没找到，遍历所有子元素查找
                                    if not rId:
                                        for elem in shape.iter():
                                            if "embed" in elem.attrib:
                                                rId = elem.get("embed")
                                                break
                                            elif "link" in elem.attrib:
                                                rId = elem.get("link")
                                                break

                                    # 如果还是没找到，尝试从父元素查找
                                    if not rId:
                                        parent = shape.getparent()
                                        if parent is not None:
                                            for elem in parent.iter():
                                                if "embed" in elem.attrib:
                                                    rId = elem.get("embed")
                                                    break
                                                elif "link" in elem.attrib:
                                                    rId = elem.get("link")
                                                    break

                                    if rId and rId in doc.part.related_parts:
                                        try:
                                            # 获取图片数据
                                            image_part = doc.part.related_parts[rId]
                                            image_data = image_part.blob

                                            # 验证图片数据
                                            if image_data and len(image_data) > 0:
                                                # 创建图片对象，自动计算尺寸
                                                img_buffer = io.BytesIO(image_data)

                                                # 尝试获取图片尺寸
                                                try:
                                                    from PIL import Image as PILImage

                                                    pil_img = PILImage.open(img_buffer)
                                                    img_width, img_height = pil_img.size

                                                    # 计算合适的PDF尺寸（保持宽高比）
                                                    max_width = 400
                                                    max_height = 300
                                                    ratio = min(max_width / img_width, max_height / img_height)
                                                    pdf_width = img_width * ratio
                                                    pdf_height = img_height * ratio

                                                    # 重置缓冲区
                                                    img_buffer.seek(0)
                                                    img = RLImage(img_buffer, width=pdf_width, height=pdf_height)
                                                except Exception:
                                                    # 如果无法获取尺寸，使用默认尺寸
                                                    img_buffer.seek(0)
                                                    img = RLImage(img_buffer, width=400, height=300)

                                                story.append(img)
                                                story.append(Spacer(1, 12))
                                                logger.info(f"成功添加图片，大小: {len(image_data)} bytes")
                                            else:
                                                logger.warning("图片数据为空")
                                                story.append(Paragraph("[图片数据为空]", normal_style))
                                                story.append(Spacer(1, 12))
                                        except Exception as img_data_error:
                                            logger.warning(f"处理图片数据时出错: {img_data_error}")
                                            story.append(Paragraph("[图片处理失败]", normal_style))
                                            story.append(Spacer(1, 12))
                                    else:
                                        logger.warning(f"无法找到图片关系ID: {rId}")
                                        story.append(Paragraph("[图片引用无效]", normal_style))
                                        story.append(Spacer(1, 12))

                                # 如果没有找到任何图片元素，尝试其他方法
                                if not image_elements:
                                    # 尝试从文档的图片集合中获取
                                    try:
                                        for rel_id, rel_part in doc.part.related_parts.items():
                                            if hasattr(rel_part, "blob") and rel_part.blob:
                                                # 检查是否是图片类型
                                                content_type = getattr(rel_part, "content_type", "")
                                                if "image" in content_type or rel_id.startswith("rId"):
                                                    try:
                                                        img_buffer = io.BytesIO(rel_part.blob)
                                                        from PIL import Image as PILImage

                                                        pil_img = PILImage.open(img_buffer)

                                                        # 计算合适的PDF尺寸
                                                        img_width, img_height = pil_img.size
                                                        max_width = 400
                                                        max_height = 300
                                                        ratio = min(max_width / img_width, max_height / img_height)
                                                        pdf_width = img_width * ratio
                                                        pdf_height = img_height * ratio

                                                        img_buffer.seek(0)
                                                        img = RLImage(img_buffer, width=pdf_width, height=pdf_height)
                                                        story.append(img)
                                                        story.append(Spacer(1, 12))
                                                        logger.info(f"从文档图片集合中找到图片: {rel_id}")
                                                        break
                                                    except Exception:
                                                        continue
                                    except Exception as e:
                                        logger.warning(f"从文档图片集合获取图片失败: {e}")

                                # 方法6: 直接从inline_shapes获取图片
                                if not image_elements and hasattr(doc, "inline_shapes") and doc.inline_shapes:
                                    try:
                                        for shape in doc.inline_shapes:
                                            if hasattr(shape, "_element"):
                                                # 尝试从inline shape获取图片
                                                shape_element = shape._element
                                                for elem in shape_element.iter():
                                                    if "embed" in elem.attrib:
                                                        rId = elem.get("embed")
                                                        if rId and rId in doc.part.related_parts:
                                                            try:
                                                                image_part = doc.part.related_parts[rId]
                                                                image_data = image_part.blob

                                                                if image_data and len(image_data) > 0:
                                                                    img_buffer = io.BytesIO(image_data)
                                                                    from PIL import Image as PILImage

                                                                    pil_img = PILImage.open(img_buffer)

                                                                    img_width, img_height = pil_img.size
                                                                    max_width = 400
                                                                    max_height = 300
                                                                    ratio = min(max_width / img_width, max_height / img_height)
                                                                    pdf_width = img_width * ratio
                                                                    pdf_height = img_height * ratio

                                                                    img_buffer.seek(0)
                                                                    img = RLImage(
                                                                        img_buffer, width=pdf_width, height=pdf_height
                                                                    )
                                                                    story.append(img)
                                                                    story.append(Spacer(1, 12))
                                                                    logger.info(f"从inline_shapes中找到图片: {rId}")
                                                                    break
                                                            except Exception as e:
                                                                logger.warning(f"处理inline shape图片失败: {e}")
                                                                continue
                                                    elif "link" in elem.attrib:
                                                        rId = elem.get("link")
                                                        if rId and rId in doc.part.related_parts:
                                                            # 处理链接图片
                                                            try:
                                                                image_part = doc.part.related_parts[rId]
                                                                image_data = image_part.blob

                                                                if image_data and len(image_data) > 0:
                                                                    img_buffer = io.BytesIO(image_data)
                                                                    from PIL import Image as PILImage

                                                                    pil_img = PILImage.open(img_buffer)

                                                                    img_width, img_height = pil_img.size
                                                                    max_width = 400
                                                                    max_height = 300
                                                                    ratio = min(max_width / img_width, max_height / img_height)
                                                                    pdf_width = img_width * ratio
                                                                    pdf_height = img_height * ratio

                                                                    img_buffer.seek(0)
                                                                    img = RLImage(
                                                                        img_buffer, width=pdf_width, height=pdf_height
                                                                    )
                                                                    story.append(img)
                                                                    story.append(Spacer(1, 12))
                                                                    logger.info(f"从inline_shapes链接中找到图片: {rId}")
                                                                    break
                                                            except Exception as e:
                                                                logger.warning(f"处理inline shape链接图片失败: {e}")
                                                                continue
                                    except Exception as e:
                                        logger.warning(f"处理inline_shapes失败: {e}")

                                # 方法7: 从文档的所有关系部分查找图片
                                if not image_elements:
                                    try:
                                        # 遍历所有关系部分，查找图片
                                        for rel_id, rel_part in doc.part.related_parts.items():
                                            if hasattr(rel_part, "blob") and rel_part.blob:
                                                content_type = getattr(rel_part, "content_type", "")
                                                # 检查是否是图片类型
                                                if any(
                                                    img_type in content_type.lower()
                                                    for img_type in ["image", "jpeg", "png", "gif", "bmp", "tiff"]
                                                ):
                                                    try:
                                                        img_buffer = io.BytesIO(rel_part.blob)
                                                        from PIL import Image as PILImage

                                                        pil_img = PILImage.open(img_buffer)

                                                        # 计算合适的PDF尺寸
                                                        img_width, img_height = pil_img.size
                                                        max_width = 400
                                                        max_height = 300
                                                        ratio = min(max_width / img_width, max_height / img_height)
                                                        pdf_width = img_width * ratio
                                                        pdf_height = img_height * ratio

                                                        img_buffer.seek(0)
                                                        img = RLImage(img_buffer, width=pdf_width, height=pdf_height)
                                                        story.append(img)
                                                        story.append(Spacer(1, 12))
                                                        logger.info(f"从关系部分找到图片: {rel_id}, 类型: {content_type}")
                                                        # 不要break，继续查找更多图片
                                                    except Exception as e:
                                                        logger.warning(f"处理关系部分图片失败: {e}")
                                                        continue
                                    except Exception as e:
                                        logger.warning(f"从关系部分查找图片失败: {e}")

                            except Exception as img_error:
                                logger.warning(f"处理图片时出错: {img_error}")
                                # 如果图片处理失败，添加占位符
                                story.append(Paragraph("[图片处理异常]", normal_style))
                                story.append(Spacer(1, 12))

                # 生成PDF
                pdf_doc.build(story)
                pdf_content = pdf_buffer.getvalue()
                pdf_buffer.close()

                # 清理临时文件
                try:
                    os.unlink(temp_docx_path)
                except Exception:
                    pass

                if len(pdf_content) == 0:
                    return False, "转换后的文件为空，可能是Word文档内容无法识别", None

                return True, pdf_content, "word_to_pdf"

            except Exception as conversion_error:
                # 清理临时文件
                try:
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
                except Exception:
                    pass

                # 提供更详细的错误信息
                error_msg = str(conversion_error)
                if "Package not found" in error_msg:
                    return False, "Word文件损坏或格式不支持，请检查文件完整性", None
                elif "Permission denied" in error_msg:
                    return False, "文件访问权限不足，请检查文件权限", None
                else:
                    return False, f"Word转PDF转换失败: {error_msg}", None

        except Exception as e:
            logger.error(f"Word转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def pdf_to_images(self, pdf_file, dpi=150):
        """PDF转图片"""
        try:
            if not FITZ_AVAILABLE:
                return False, "PyMuPDF未安装，无法进行PDF转换", None

            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            images = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                mat = fitz.Matrix(dpi / 72, dpi / 72)  # 设置DPI
                pix = page.get_pixmap(matrix=mat)

                # 转换为PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))

                # 转换为base64
                img_buffer = io.BytesIO()
                img.save(img_buffer, format="PNG")
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode()

                images.append({"page": page_num + 1, "data": img_base64, "width": img.width, "height": img.height})

            doc.close()
            return True, images, "pdf_to_images"

        except Exception as e:
            logger.error(f"PDF转图片失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def images_to_pdf(self, image_files):
        """图片转PDF"""
        try:
            images = []

            for img_file in image_files:
                img = Image.open(img_file)
                if img.mode != "RGB":
                    img = img.convert("RGB")
                images.append(img)

            if not images:
                return False, "没有有效的图片文件", None

            # 创建PDF
            pdf_buffer = io.BytesIO()
            if len(images) == 1:
                images[0].save(pdf_buffer, format="PDF")
            else:
                images[0].save(pdf_buffer, format="PDF", save_all=True, append_images=images[1:])

            pdf_content = pdf_buffer.getvalue()
            return True, pdf_content, "images_to_pdf"

        except Exception as e:
            logger.error(f"图片转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def text_to_pdf(self, text_content):
        """文本转PDF"""
        try:
            # 检查reportlab库是否可用
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            except ImportError:
                return False, "reportlab库未安装，无法进行文本转PDF转换", None

            # 创建PDF缓冲区
            pdf_buffer = io.BytesIO()

            # 创建PDF文档
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            story = []

            # 获取样式
            styles = getSampleStyleSheet()
            normal_style = styles["Normal"]

            # 设置中文字体支持
            try:
                # 使用reportlab内置的中文字体
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont

                pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                normal_style.fontName = "STSong-Light"
                normal_style.fontSize = 12
                normal_style.leading = 14
                normal_style.alignment = 0  # 左对齐
            except Exception:
                try:
                    # 尝试使用系统中文字体
                    import platform

                    if platform.system() == "Darwin":  # macOS
                        try:
                            pdfmetrics.registerFont(TTFont("PingFang", "/System/Library/Fonts/PingFang.ttc"))
                            normal_style.fontName = "PingFang"
                        except Exception:
                            pdfmetrics.registerFont(TTFont("HiraginoSans", "/System/Library/Fonts/STHeiti Light.ttc"))
                            normal_style.fontName = "HiraginoSans"
                    elif platform.system() == "Windows":
                        pdfmetrics.registerFont(TTFont("SimSun", "C:/Windows/Fonts/simsun.ttc"))
                        normal_style.fontName = "SimSun"
                    else:  # Linux
                        pdfmetrics.registerFont(TTFont("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
                        normal_style.fontName = "DejaVuSans"

                    normal_style.fontSize = 12
                    normal_style.leading = 14
                    normal_style.alignment = 0
                except Exception as e2:
                    # 如果都不可用，使用默认字体
                    normal_style.fontName = "Helvetica"
                    normal_style.fontSize = 12
                    normal_style.leading = 14
                    normal_style.alignment = 0
                    logger.warning(f"无法加载中文字体，将使用默认字体: {str(e2)}")

            # 处理文本内容
            lines = text_content.split("\n")
            for line in lines:
                line = line.strip()
                if line:
                    # 创建段落
                    paragraph = Paragraph(line, normal_style)
                    story.append(paragraph)
                    story.append(Spacer(1, 6))  # 添加间距
                else:
                    # 空行
                    story.append(Spacer(1, 12))

            # 生成PDF
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()

            if len(pdf_content) == 0:
                return False, "生成的PDF文件为空", None

            return True, pdf_content, "text_to_pdf"

        except Exception as e:
            logger.error(f"文本转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def pdf_to_text(self, pdf_file):
        """PDF转文本"""
        try:
            # 优先使用PyMuPDF，如果不可用则使用pypdf
            if FITZ_AVAILABLE:
                return self._pdf_to_text_fitz(pdf_file)
            elif PYPDF_AVAILABLE:
                return self._pdf_to_text_pypdf(pdf_file)
            else:
                return False, "PDF处理库未安装，无法进行PDF转文本转换", None

            # 创建临时文件
            import os
            import tempfile

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                # 写入PDF内容
                for chunk in pdf_file.chunks():
                    temp_pdf.write(chunk)
                temp_pdf_path = temp_pdf.name

            try:
                # 打开PDF文件
                doc = fitz.open(temp_pdf_path)

                # 检查PDF是否为空或损坏
                if len(doc) == 0:
                    doc.close()
                    return False, "PDF文件为空或损坏", None

                text_content = ""
                total_pages = len(doc)

                # 逐页提取文本
                for page_num in range(total_pages):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        text_content += page_text
                        if page_num < total_pages - 1:
                            text_content += "\n\n"  # 页面间添加空行
                    except Exception as page_error:
                        logger.warning(f"提取第{page_num + 1}页文本时出错: {str(page_error)}")
                        text_content += f"\n[第{page_num + 1}页文本提取失败]\n"

                doc.close()

                # 检查提取的文本内容
                if not text_content.strip():
                    return False, "PDF文件不包含可提取的文本内容（可能是扫描版PDF，建议使用OCR工具）", None

                # 如果文本内容很少，可能是扫描版PDF
                if len(text_content.strip()) < 10:
                    return False, "提取的文本内容过少，可能是扫描版PDF，建议使用OCR工具", None

                return True, text_content, "pdf_to_text"

            finally:
                # 清理临时文件
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)

        except Exception as e:
            logger.error(f"PDF转文本失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def txt_to_pdf(self, txt_file):
        """TXT文件转PDF"""
        try:
            # 智能编码检测和读取txt文件内容
            raw_content = txt_file.read()

            # 尝试多种编码方式
            encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
            txt_content = None

            for encoding in encodings:
                try:
                    txt_content = raw_content.decode(encoding)
                    logger.info(f"成功使用 {encoding} 编码读取文件")
                    break
                except UnicodeDecodeError:
                    continue

            if txt_content is None:
                # 如果所有编码都失败，使用错误替换模式
                txt_content = raw_content.decode("utf-8", errors="replace")
                logger.warning("使用utf-8错误替换模式读取文件")

            # 调用text_to_pdf方法
            return self.text_to_pdf(txt_content)

        except Exception as e:
            logger.error(f"TXT文件转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def _pdf_to_text_fitz(self, pdf_file):
        """使用PyMuPDF进行PDF转文本"""
        # 创建临时文件
        import os
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            # 写入PDF内容
            for chunk in pdf_file.chunks():
                temp_pdf.write(chunk)
            temp_pdf_path = temp_pdf.name

        try:
            # 打开PDF文件
            doc = fitz.open(temp_pdf_path)

            # 检查PDF是否为空或损坏
            if len(doc) == 0:
                doc.close()
                return False, "PDF文件为空或损坏", None

            text_content = ""
            total_pages = len(doc)

            # 逐页提取文本
            for page_num in range(total_pages):
                try:
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    text_content += page_text
                    if page_num < total_pages - 1:
                        text_content += "\n\n"  # 页面间添加空行
                except Exception as page_error:
                    logger.warning(f"提取第{page_num + 1}页文本时出错: {str(page_error)}")
                    text_content += f"\n[第{page_num + 1}页文本提取失败]\n"

            doc.close()

            # 检查提取的文本内容
            if not text_content.strip():
                return False, "PDF文件不包含可提取的文本内容（可能是扫描版PDF，建议使用OCR工具）", None

            # 如果文本内容很少，可能是扫描版PDF
            if len(text_content.strip()) < 10:
                return False, "提取的文本内容过少，可能是扫描版PDF，建议使用OCR工具", None

            return True, text_content, "pdf_to_text"

        finally:
            # 清理临时文件
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)

    def _pdf_to_text_pypdf(self, pdf_file):
        """使用pypdf进行PDF转文本"""
        try:
            # 重置文件指针
            pdf_file.seek(0)

            # 使用pypdf读取PDF
            pdf_reader = pypdf.PdfReader(pdf_file)

            # 检查PDF是否为空或损坏
            if len(pdf_reader.pages) == 0:
                return False, "PDF文件为空或损坏", None

            text_content = ""
            total_pages = len(pdf_reader.pages)

            # 逐页提取文本
            for page_num in range(total_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text_content += page_text
                    if page_num < total_pages - 1:
                        text_content += "\n\n"  # 页面间添加空行
                except Exception as page_error:
                    logger.warning(f"提取第{page_num + 1}页文本时出错: {str(page_error)}")
                    text_content += f"\n[第{page_num + 1}页文本提取失败]\n"

            # 检查提取的文本内容
            if not text_content.strip():
                return False, "PDF文件不包含可提取的文本内容（可能是扫描版PDF，建议使用OCR工具）", None

            # 如果文本内容很少，可能是扫描版PDF
            if len(text_content.strip()) < 10:
                return False, "提取的文本内容过少，可能是扫描版PDF，建议使用OCR工具", None

            return True, text_content, "pdf_to_text"

        except Exception as e:
            logger.error(f"pypdf PDF转文本失败: {str(e)}")
            return False, f"pypdf PDF转文本失败: {str(e)}", None


@csrf_exempt
@require_http_methods(["POST"])
def pdf_converter_api(request):
    """PDF转换API主入口"""
    try:
        # 导入模型
        import time

        from .models.legacy_models import PDFConversionRecord

        # 创建转换器实例
        converter = PDFConverter()

        # 添加调试信息
        logger.info(f"PDF转换API请求: POST数据={dict(request.POST)}, FILES={list(request.FILES.keys())}")

        # 支持JSON和表单数据
        if request.content_type and "application/json" in request.content_type:
            try:
                data = json.loads(request.body)
                conversion_type = data.get("type", "")
                text_content = data.get("text_content", "")
                logger.info(
                    f"JSON数据: type={conversion_type}, text_content_length={len(text_content) if text_content else 0}"
                )
            except json.JSONDecodeError:
                return JsonResponse({"success": False, "error": "无效的JSON数据"}, status=400)
        else:
            # 表单数据
            conversion_type = request.POST.get("type", "")
            text_content = request.POST.get("text_content", "")
            logger.info(f"表单数据: type={conversion_type}, text_content_length={len(text_content) if text_content else 0}")

        # 添加更详细的调试信息
        if "file" in request.FILES:
            file = request.FILES["file"]
            logger.info(f"上传文件信息: 名称={file.name}, 大小={file.size}, 类型={file.content_type}")

        conversion_type = conversion_type or ""

        # 检查是否有文件上传（文本转PDF除外）
        if conversion_type != "text-to-pdf":
            if "file" not in request.FILES:
                logger.warning("PDF转换API: 没有上传文件")
                return JsonResponse({"success": False, "error": "没有上传文件"}, status=400)
            file = request.FILES["file"]
        else:
            # 文本转PDF不需要文件上传
            file = None

        # 创建转换记录（如果用户已登录）
        conversion_record = None
        if request.user.is_authenticated:
            if conversion_type == "text-to-pdf":
                # 文本转PDF的特殊处理
                text_content = request.POST.get("text_content", "")
                original_filename = f"text_content_{len(text_content)}_chars.txt"
                file_size = len(text_content.encode("utf-8"))
            else:
                original_filename = file.name
                file_size = file.size

            conversion_record = PDFConversionRecord.objects.create(
                user=request.user,
                conversion_type=conversion_type.replace("-", "_"),
                original_filename=original_filename,
                file_size=file_size,
                status="processing",
            )

        start_time = time.time()

        # 验证转换类型
        valid_types = [
            "pdf-to-word",
            "word-to-pdf",
            "pdf-to-image",
            "image-to-pdf",
            "text-to-pdf",
            "pdf-to-text",
            "txt-to-pdf",
        ]
        if conversion_type not in valid_types:
            return JsonResponse({"success": False, "error": f"不支持的转换类型: {conversion_type}"}, status=400)

        # 根据转换类型验证文件格式
        if conversion_type == "text-to-pdf":
            # 文本转PDF不需要文件验证
            # text_content已经在上面从JSON或表单数据中获取
            if not text_content.strip():
                is_valid, message = False, "请输入要转换的文本内容"
            else:
                is_valid, message = True, "文本内容验证通过"
        elif conversion_type == "pdf-to-word":
            is_valid, message = converter.validate_file(file, "pdf")
        elif conversion_type == "word-to-pdf":
            is_valid, message = converter.validate_file(file, "word")
        elif conversion_type == "pdf-to-image":
            is_valid, message = converter.validate_file(file, "pdf")
        elif conversion_type == "image-to-pdf":
            is_valid, message = converter.validate_file(file, "image")
        elif conversion_type == "pdf-to-text":
            is_valid, message = converter.validate_file(file, "pdf")
        elif conversion_type == "txt-to-pdf":
            is_valid, message = converter.validate_file(file, "text")
        else:
            is_valid, message = True, "文件验证通过"

        if not is_valid:
            # 检查是否包含转换建议
            if "建议的转换类型" in message:
                # 解析建议的转换类型
                suggested_types = []
                lines = message.split("\n")
                for line in lines:
                    if line.strip().startswith("•"):
                        # 提取转换类型
                        conv_type = line.split("(")[1].split(")")[0] if "(" in line else None
                        if conv_type:
                            suggested_types.append(conv_type)

                return JsonResponse(
                    {"success": False, "error": message, "suggested_types": suggested_types, "needs_type_switch": True},
                    status=400,
                )
            else:
                return JsonResponse({"success": False, "error": message}, status=400)

        # 执行转换
        if conversion_type == "text-to-pdf":
            # text_content已经在上面从JSON或表单数据中获取
            success, result, file_type = converter.text_to_pdf(text_content)
        elif conversion_type == "pdf-to-word":
            success, result, file_type = converter.pdf_to_word(file)
        elif conversion_type == "word-to-pdf":
            success, result, file_type = converter.word_to_pdf(file)
        elif conversion_type == "pdf-to-image":
            success, result, file_type = converter.pdf_to_images(file)
        elif conversion_type == "image-to-pdf":
            success, result, file_type = converter.images_to_pdf([file])
        elif conversion_type == "pdf-to-text":
            success, result, file_type = converter.pdf_to_text(file)
        elif conversion_type == "txt-to-pdf":
            success, result, file_type = converter.txt_to_pdf(file)
        else:
            return JsonResponse({"success": False, "error": "未知的转换类型"}, status=400)

        # 计算转换时间
        conversion_time = time.time() - start_time

        if not success:
            # 更新转换记录为失败状态（如果存在）
            if conversion_record:
                conversion_record.status = "failed"
                conversion_record.error_message = result
                conversion_record.conversion_time = conversion_time
                conversion_record.save()

            return JsonResponse({"success": False, "error": result}, status=500)

        # 保存转换结果
        output_filename = f"{uuid.uuid4()}_{conversion_type.replace('-', '_')}"

        if file_type == "pdf_to_word":
            output_filename += ".docx"
            default_storage.save(f"converted/{output_filename}", ContentFile(result))
            # 设置下载链接
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
        elif file_type == "word_to_pdf":
            output_filename += ".pdf"
            default_storage.save(f"converted/{output_filename}", ContentFile(result))
            # 设置下载链接
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
        elif file_type == "pdf_to_images":
            # 创建ZIP文件包含所有图片
            import base64
            import zipfile
            from io import BytesIO

            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for i, img_data in enumerate(result):
                    # 解码base64图片数据
                    img_bytes = base64.b64decode(img_data["data"])
                    # 添加到ZIP文件
                    zip_file.writestr(f"page_{i+1}.png", img_bytes)

            zip_content = zip_buffer.getvalue()
            zip_buffer.close()

            # 保存ZIP文件
            output_filename += "_images.zip"
            default_storage.save(f"converted/{output_filename}", ContentFile(zip_content))
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()

            # 返回下载链接
            return JsonResponse(
                {
                    "success": True,
                    "type": "file",
                    "download_url": download_url,
                    "filename": output_filename,
                    "original_filename": file.name,
                    "file_size": len(zip_content),
                    "total_pages": len(result),
                    "message": f"已转换{len(result)}页，打包为ZIP文件供下载",
                    "conversion_type": conversion_type,
                }
            )
        elif file_type == "images_to_pdf":
            output_filename += ".pdf"
            default_storage.save(f"converted/{output_filename}", ContentFile(result))
            # 设置下载链接
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
        elif file_type == "text_to_pdf":
            output_filename += ".pdf"
            default_storage.save(f"converted/{output_filename}", ContentFile(result))
            # 设置下载链接
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
        elif file_type == "pdf_to_text":
            output_filename += ".txt"
            default_storage.save(f"converted/{output_filename}", ContentFile(result.encode("utf-8")))
            # 设置下载链接
            download_url = f"/tools/api/pdf-converter/download/{output_filename}/"

            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = "success"
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
        else:
            return JsonResponse({"success": False, "error": "未知的文件类型"}, status=500)

        # 确定原始文件名
        if conversion_type == "text-to-pdf":
            original_filename = "文本内容"
        elif conversion_type == "pdf-to-text":
            original_filename = file.name if file else "PDF文件"
        else:
            original_filename = file.name if file else "文件"

        return JsonResponse(
            {
                "success": True,
                "type": "file",
                "download_url": download_url,
                "filename": output_filename,
                "original_filename": original_filename,
                "conversion_type": conversion_type,
            }
        )

    except Exception as e:
        logger.error(f"PDF转换测试API错误: {str(e)}")
        return JsonResponse({"success": False, "error": f"服务器错误: {str(e)}"}, status=500)
